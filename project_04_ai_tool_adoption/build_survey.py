"""
Generates the pre/post usability survey for Project 4.
Plain Word aesthetic — numbered questions, response options as indented text,
section breaks with simple rules. Looks like a Qualtrics export or a
well-formatted Google Form printout, not a designed artifact.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BBBBBB')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    font(run, size=12, bold=True)

def instruction(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    font(run, size=10, italic=True, color=(100, 100, 100))

def q(num, text, sub=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    rn = p.add_run(f'{num}.  ')
    font(rn, size=11, bold=True)
    rq = p.add_run(text)
    font(rq, size=11)
    if sub:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent  = Inches(0.3)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(3)
        rs = p2.add_run(sub)
        font(rs, size=10, italic=True, color=(110, 110, 110))

def scale_row(label_left, label_right, points=5):
    """Renders a labeled Likert row: label | 1  2  3  4  5 | label"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    rl = p.add_run(f'{label_left:<22}')
    font(rl, size=10, color=(80, 80, 80))
    for i in range(1, points + 1):
        rb = p.add_run(f'  □ {i}  ')
        font(rb, size=10)
    rr = p.add_run(f'  {label_right}')
    font(rr, size=10, color=(80, 80, 80))

def options(opts, other=False):
    for opt in opts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.35)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(f'□  {opt}')
        font(run, size=11)
    if other:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.35)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run('□  Other (please specify): ___________________________')
        font(run, size=11)

def open_end(lines=3):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run('_' * 72)
        font(run, size=10, color=(180, 180, 180))

# ═══════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('AI Research Platform: Analyst Experience Survey')
font(run, size=15, bold=True)

p2 = doc.add_paragraph()
run2 = p2.add_run(
    'Pre- and Post-Pilot Survey  |  AI Tool Adoption Study  |  '
    'Estimated time: 10–12 minutes'
)
font(run2, size=11, color=(90, 90, 90))

p3 = doc.add_paragraph()
run3 = p3.add_run(
    'Your responses are confidential. Individual answers will not be shared '
    'with managers or leadership. Results are reported in aggregate only.'
)
font(run3, size=10, italic=True, color=(120, 120, 120))

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION A: SCREENER / BACKGROUND
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section A — Background')
instruction(
    'These questions help us understand your starting point before the pilot. '
    'There are no right or wrong answers.'
)

q(1, 'Which team are you on?')
options(['Customer Insights', 'Product Research', 'Data Science'])

q(2, 'How long have you been in your current role?')
options(['Less than 1 year', '1–2 years', '3–5 years', 'More than 5 years'])

q(3, 'Before this pilot, how would you describe your experience with AI-assisted '
      'research tools (e.g., automated coding, LLM-assisted analysis)?',
   sub='Select one.')
options([
    'No prior experience',
    'Experimented with tools informally (personal projects, demos)',
    'Used AI tools occasionally in my work',
    'Used AI tools regularly as part of my workflow',
])

q(4, 'How do you currently handle qualitative analysis tasks — transcript coding, '
      'survey open-end processing, theme identification?',
   sub='Select all that apply.')
options([
    'Manually, in Word or Excel',
    'Using dedicated qualitative software (NVivo, Atlas.ti, Dedoose)',
    'Using a combination of tools and manual review',
    'Delegating to a junior analyst or contractor',
], other=True)

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION B: USABILITY (SUS-derived, pre and post)
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section B — Usability')
instruction(
    'Rate each statement on a 1–5 scale where 1 = Strongly Disagree '
    'and 5 = Strongly Agree. Answer based on your experience with the '
    'platform so far. If you have not used it yet, skip to Section C.'
)

usability_items = [
    ('B1', 'I think I would use this platform frequently.'),
    ('B2', 'I found the platform unnecessarily complex.'),
    ('B3', 'I felt confident using the platform without needing outside help.'),
    ('B4', 'I thought there was too much inconsistency in this platform.'),
    ('B5', 'I would imagine most people would learn to use this platform quickly.'),
    ('B6', 'I found the platform very cumbersome to use.'),
    ('B7', 'I felt very confident using the platform.'),
    ('B8', 'I needed to learn a lot of things before I could get going with this platform.'),
]

for num, text in usability_items:
    q(num, text)
    scale_row('Strongly disagree', 'Strongly agree', points=5)

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION C: TRUST CALIBRATION
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section C — Trust and Output Quality')
instruction(
    'These questions focus on how you evaluate and respond to the platform\'s outputs. '
    'Trust calibration — knowing when to trust and when to verify — is one of the '
    'strongest predictors of effective tool use in our research.'
)

q('C1', 'When the platform produces a coded or summarized output, how often do you '
         'verify it against the original source material?',
   sub='Select one.')
options([
    'Always — I check every output before using it',
    'Usually — I spot-check a portion of the output',
    'Sometimes — only when something looks off',
    'Rarely — I generally trust the output as-is',
    'I have not used the platform for this type of task',
])

q('C2', 'Has the platform ever given you an output that was incorrect or misleading?')
options([
    'Yes — and I caught it before it affected my work',
    'Yes — and I did not catch it immediately',
    'No, not that I noticed',
    'I am not sure',
])

q('C3', 'When you encounter an error or unexpected output from the platform, '
         'what do you typically do?',
   sub='Select all that apply.')
options([
    'Re-run the query with different inputs',
    'Flag it using the in-platform error reporting feature',
    'Document it in my own notes',
    'Tell a colleague',
    'Move on without doing anything',
], other=True)

q('C4', 'On a scale of 1–10, how accurately calibrated do you feel your trust '
         'in the platform is right now?',
   sub='1 = I distrust it almost entirely  |  10 = I trust it almost entirely without checking')
scale_row('Distrust entirely', 'Trust entirely', points=10)

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION D: WORKFLOW INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section D — Workflow Integration')
instruction(
    'These questions ask about how the platform fits — or does not fit — '
    'into your actual day-to-day work.'
)

q('D1', 'On average, how many times per week do you use the platform for a '
         'real work task (not a test or training session)?',
   sub='Select one.')
options([
    'I have not used it for a real task yet',
    'Less than once a week',
    '1–2 times per week',
    '3–5 times per week',
    'More than 5 times per week',
])

q('D2', 'For which types of tasks have you used the platform?',
   sub='Select all that apply.')
options([
    'Transcript coding or thematic analysis',
    'Survey open-end processing or summarization',
    'Literature or document review',
    'Research brief or summary generation',
    'Data labeling or categorization',
], other=True)

q('D3', 'How much effort does it currently take to incorporate the platform\'s '
         'outputs into your existing workflow?',
   sub='1 = Almost no effort  |  10 = Significant effort (copy-paste, reformatting, manual verification)')
scale_row('Almost no effort', 'Significant effort', points=10)

q('D4', 'What is the biggest friction point when using the platform in your work?')
open_end(lines=3)

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION E: ONBOARDING REFLECTION
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section E — Onboarding')
instruction(
    'These questions help us improve the training and onboarding process '
    'for future rollouts.'
)

q('E1', 'How would you rate the quality of the initial onboarding session?',
   sub='1 = Not useful at all  |  5 = Extremely useful')
scale_row('Not useful', 'Extremely useful', points=5)

q('E2', 'After the onboarding session, did you feel prepared to use the platform '
         'independently on a real project?')
options([
    'Yes, fully prepared',
    'Mostly prepared, with some gaps',
    'Partially prepared — I had significant gaps',
    'Not prepared',
])

q('E3', 'What scenario or use case do you wish had been covered in onboarding?')
open_end(lines=3)

q('E4', 'Have you gone back to the onboarding materials (recordings, documentation, '
         'guides) after the initial session?')
options([
    'Yes, multiple times',
    'Yes, once',
    'No, but I would if I knew where to find them',
    'No — I did not find them helpful the first time',
    'No — I have not needed them',
])

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION F: OPEN-END CLOSE
# ═══════════════════════════════════════════════════════════════════════════
section_header('Section F — Anything Else?')
instruction(
    'Two final questions. These are optional but often produce the most useful feedback.'
)

q('F1', 'If you could change one thing about the platform or how it was introduced '
         'to your team, what would it be?')
open_end(lines=4)

q('F2', 'Is there anything about your experience with the platform that we have '
         'not asked about and that you think is important for us to understand?')
open_end(lines=4)

p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(12)
r_end = p_end.add_run(
    'Thank you for completing this survey. Your feedback directly informs '
    'improvements to the platform and the next onboarding cycle.'
)
font(r_end, size=10, italic=True, color=(100, 100, 100))

out = 'survey_ai_tool_adoption.docx'
doc.save(out)
print(f'Saved: {out}')
