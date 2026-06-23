"""
Builds all three portfolio showcase documents as .docx Word files.
Run:  python3 build_word_docs.py
Outputs:
  01_survey_research/Survey Methods Showcase.docx
  05_qualitative_methods/Interview Methods Showcase.docx
  06_experiment_design/Experiment Design Showcase.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = Path("/Users/jbarajas/Documents/Leaving RAND/Job Hunting/Upwork Portfolio")

# ── Colors ────────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x2D, 0x6A, 0x9F)
ORANGE = RGBColor(0xE8, 0x77, 0x22)
GREEN  = RGBColor(0x2E, 0x7D, 0x52)
GRAY   = RGBColor(0x55, 0x55, 0x55)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLUE_LIGHT_HEX  = "E8F2FB"
ORANGE_LIGHT_HEX = "FEF3E8"
GREEN_LIGHT_HEX  = "E8F5EE"
TABLE_HEADER_HEX = "2D6A9F"
TABLE_ROW1_HEX   = "E8F2FB"
TABLE_ROW2_HEX   = "FFFFFF"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), kwargs.get("val", "single"))
        border.set(qn("w:sz"), kwargs.get("sz", "4"))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), kwargs.get("color", "D0D8E4"))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def para(doc, text, bold=False, size=11, color=None, align=None, space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def heading(doc, text, level=1, color=BLUE, size=None):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.text = text
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.color.rgb = color
    return p


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D8E4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def callout_table(doc, head_text, body_text, bg_hex=BLUE_LIGHT_HEX, head_color=None):
    if head_color is None:
        head_color = BLUE
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(1.1)
    tbl.columns[1].width = Inches(5.2)
    hc = tbl.cell(0, 0)
    bc = tbl.cell(0, 1)
    set_cell_bg(hc, bg_hex)
    set_cell_bg(bc, bg_hex)
    hp = hc.paragraphs[0]
    hr = hp.add_run(head_text)
    hr.bold = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = head_color
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    br = bp.add_run(body_text)
    br.font.size = Pt(9)
    br.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def method_header(doc, num, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    badge = p.add_run(f"METHOD {num:02d}   ")
    badge.bold = True
    badge.font.size = Pt(9)
    badge.font.color.rgb = ORANGE
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = BLUE
    return p


def design_header(doc, num, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.clear()
    badge = p.add_run(f"DESIGN {num:02d}   ")
    badge.bold = True
    badge.font.size = Pt(9)
    badge.font.color.rgb = ORANGE
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = BLUE
    return p


def spec_table(doc, rows, col_widths=(1.5, 5.0)):
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows):
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        c0.width = Inches(col_widths[0])
        c1.width = Inches(col_widths[1])
        bg = TABLE_HEADER_HEX if i == 0 else (TABLE_ROW1_HEX if i % 2 == 1 else TABLE_ROW2_HEX)
        set_cell_bg(c0, bg)
        set_cell_bg(c1, bg)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = WHITE if i == 0 else DARK
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(9)
        r1.font.color.rgb = WHITE if i == 0 else DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def option_list(doc, options):
    for opt in options:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(opt)
        run.font.size = Pt(10)


def probe_list(doc, probes):
    for probe in probes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'"{probe}"')
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = DARK


def q_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ORANGE


def footer_text():
    return "Jeremy Barajas  |  barajas@alumni.usc.edu  |  RAND Corporation  |  M.B.D.S., University of Pennsylvania"


def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT 1 — Survey Methods Showcase
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def build_survey_doc():
    doc = Document()
    set_doc_margins(doc)

    # Cover
    para(doc, "SURVEY RESEARCH METHODS SHOWCASE", bold=True, size=22, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Seven question design techniques — with examples, rationale, and client value.",
         size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Jeremy Barajas  |  Behavioral Scientist & Mixed-Methods Researcher",
         bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "RAND Corporation  ·  University of Pennsylvania (M.B.D.S.)  ·  USC (B.A. Psychology, Honors)",
         size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    divider(doc)
    para(doc,
         "This document demonstrates 7 survey question formats inspired by my work at RAND Corporation "
         "and in consulting. Each example includes a design note explaining why the format was selected "
         "and what it enables analytically.",
         size=10, color=DARK, space_after=10)

    # TOC
    q_label(doc, "CONTENTS")
    for item in [
        "Method 01  —  Adaptive Branching Logic",
        "Method 02  —  Likert Matrix",
        "Method 03  —  Constant-Sum (Point Allocation)",
        "Method 04  —  Vignette / Scenario-Based Item",
        "Method 05  —  Semantic Differential Scale",
        "Method 06  —  Social Circle / Network Influence",
        "Method 07  —  Open-Ended with Structured Probe",
    ]:
        para(doc, item, size=10, space_after=2)

    doc.add_page_break()

    # ── Method 01 ─────────────────────────────────────────────────────────────
    method_header(doc, 1, "Adaptive Branching Logic")
    divider(doc)
    para(doc,
         "Branching logic routes respondents through different question paths based on prior answers, "
         "so they only see items relevant to their situation. Done well, it reduces burden and improves "
         "data quality by eliminating inapplicable items that confuse respondents.",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTION")
    para(doc, "Q1.  Have you used the AI analytics platform in the past 30 days?", size=10.5, space_after=3)
    option_list(doc, ["Yes", "No", "I was not aware this tool existed"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("→  If YES: continue to Section B (Usage & Experience)\n→  If NO or Unaware: skip to Section C (Barriers & Readiness)")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = ORANGE

    q_label(doc, "BRANCHING FLOW")
    spec_table(doc, [
        ("Q1 Response", "Next Section"),
        ("Yes (used platform)", "Section B: Usage frequency, task types, satisfaction ratings"),
        ("No / Unaware", "Section C: Primary barrier, awareness source, readiness to try"),
    ])

    callout_table(doc, "WHY THIS METHOD",
        "Branching eliminates the 'not applicable' problem — respondents who've never used a tool cannot "
        "meaningfully rate its usability. Forcing them to answer wastes their time and introduces noise. "
        "Branching also lets you use one instrument to measure both adopters and non-adopters.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "You get clean, segment-specific data with one instrument. Adoption barriers can be analyzed "
        "separately from usage patterns. Reports can say 'of the 38% who haven't adopted, here is what's "
        "stopping them' — not 'here is the average of everyone mixed together.'",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 02 ─────────────────────────────────────────────────────────────
    method_header(doc, 2, "Likert Matrix")
    divider(doc)
    para(doc,
         "A matrix groups multiple attitude items sharing the same response scale into a compact grid. "
         "It reduces survey length and cognitive switching costs — but only when items are genuinely related. "
         "A poorly constructed matrix creates straightlining (respondents checking the same column for every row).",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTION")
    para(doc, "Q2.  How much do you agree with each of the following statements about the AI platform?", size=10.5, space_after=4)

    tbl = doc.add_table(rows=6, cols=6)
    tbl.style = "Table Grid"
    headers = ["", "Strongly Disagree", "Disagree", "Neutral", "Agree", "Strongly Agree"]
    items = [
        "I understand how to use the platform effectively.",
        "The platform produces outputs I can trust without verifying.",
        "Using the platform makes me look more capable to colleagues.",
        "I worry that AI tools will reduce the value of my expertise.",
        "My organization has given me adequate training to use the platform.",
    ]
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, TABLE_HEADER_HEX)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, item in enumerate(items):
        bg = TABLE_ROW1_HEX if i % 2 == 0 else TABLE_ROW2_HEX
        tbl.cell(i+1, 0).width = Inches(2.8)
        set_cell_bg(tbl.cell(i+1, 0), bg)
        run = tbl.cell(i+1, 0).paragraphs[0].add_run(item)
        run.font.size = Pt(8.5)
        for j in range(1, 6):
            set_cell_bg(tbl.cell(i+1, j), bg)
            p = tbl.cell(i+1, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("◯").font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    callout_table(doc, "WHY THIS METHOD",
        "Five distinct constructs — self-efficacy, trust, social image, threat perception, and organizational "
        "support — are measured efficiently in one grid. Each maps to a different behavioral science lever. "
        "Items are ordered to move from external behavior to internal beliefs to reduce anchoring effects.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Pinpoints exactly which belief is blocking adoption. If trust scores are low but self-efficacy is "
        "high, the intervention is communications, not training. If self-efficacy is low and trust is high, "
        "it is the opposite. These require completely different responses and investment levels.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 03 ─────────────────────────────────────────────────────────────
    method_header(doc, 3, "Constant-Sum (Point Allocation)")
    divider(doc)
    para(doc,
         "A constant-sum item asks respondents to distribute a fixed number of points across categories. "
         "Unlike Likert scales, it forces trade-offs — saying you spend more time on one task necessarily "
         "means less on another. This produces ratio-level data about relative weight, not just ordinal rankings.",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTION")
    para(doc,
         "Q3.  Think about a typical work week. Distribute 100 points across the task types below "
         "to reflect how you currently spend your analytical work time. (Points must sum to exactly 100.)",
         size=10.5, space_after=4)

    cs_rows = [
        ("Task Category", "Points Allocated\n(must sum to 100)", "Why this matters"),
        ("Qualitative analysis (interviews, transcripts, thematic coding)", "______", "Captures cognitive burden of manual interpretation work"),
        ("Quantitative analysis (statistical modeling, data cleaning, R/Python)", "______", "Core knowledge-work output — most sensitive to AI uplift"),
        ("Writing, editing, and communications", "______", "Time recovered here signals adoption quality, not just frequency"),
        ("Administrative / project management tasks", "______", "Often hidden labor; underreported in time-diary studies"),
        ("Other", "______", "Catches residual variance; flags roles with unique workflows"),
        ("TOTAL", "100", ""),
    ]
    ct = doc.add_table(rows=len(cs_rows), cols=3)
    ct.style = "Table Grid"
    col_ws = [2.8, 1.2, 2.5]
    for i, (c0, c1, c2) in enumerate(cs_rows):
        bg = TABLE_HEADER_HEX if i == 0 else ("D4EDE4" if i == len(cs_rows)-1 else (TABLE_ROW1_HEX if i%2==1 else TABLE_ROW2_HEX))
        for j, (cell_text, cw) in enumerate(zip([c0, c1, c2], col_ws)):
            cell = ct.cell(i, j)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(8.5)
            run.bold = (i == 0 or i == len(cs_rows)-1)
            run.font.color.rgb = WHITE if i == 0 else DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    callout_table(doc, "WHY THIS METHOD",
        "Standard Likert scales let respondents rate everything as 'very important,' which reveals nothing "
        "about priority. A constant-sum forces honest prioritization. In AI adoption research, this tells "
        "you where the time is actually going — and where AI delivers the most efficiency gain per user segment.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Enables ROI modeling by role: if quantitative analysts allocate 40% to data cleaning and a tool "
        "reduces that to 10%, you can estimate recovered hours per analyst per week. Surfaces the actual "
        "cost of administrative burden on knowledge-work output.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 04 ─────────────────────────────────────────────────────────────
    method_header(doc, 4, "Vignette / Scenario-Based Item")
    divider(doc)
    para(doc,
         "A vignette presents a brief, realistic scenario and asks respondents to react to it. "
         "This is preferable to abstract attitude questions when the construct is sensitive or hard to "
         "self-report accurately. Vignettes also allow systematic variation across conditions to isolate "
         "specific factors (e.g., who produced the work, what kind of output it is).",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "SCENARIO")
    scenario_tbl = doc.add_table(rows=1, cols=1)
    scenario_tbl.style = "Table Grid"
    set_cell_bg(scenario_tbl.cell(0, 0), "F0F4F9")
    sp = scenario_tbl.cell(0, 0).paragraphs[0]
    sr = sp.add_run(
        "Your colleague sends you a 10-page research brief on a policy topic you've been asked to brief "
        "leadership on. The brief is well-organized, clearly written, and cites relevant sources. You later "
        "learn it was drafted entirely by an AI tool, then reviewed and lightly edited by your colleague before sending."
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    q_label(doc, "EXAMPLE QUESTIONS")
    para(doc,
         "Q4a.  How confident would you be using this brief to prepare your leadership presentation?",
         size=10.5, space_after=3)
    option_list(doc, [
        "I would use it with no additional verification.",
        "I would use it after spot-checking the key claims.",
        "I would not use it without a full independent review.",
    ])
    para(doc,
         "Q4b.  Would your answer change if the brief had been drafted by a junior analyst instead of an AI?",
         size=10.5, space_after=3)
    option_list(doc, ["Yes — I would be more confident.", "Yes — I would be less confident.", "No difference."])

    callout_table(doc, "WHY THIS METHOD",
        "Asking 'do you trust AI outputs?' produces socially desirable responses — people say what sounds "
        "reasonable, not what they'd actually do. The vignette bypasses this by grounding the question in "
        "a specific, plausible situation. You can also ask a reference group version: 'What would your "
        "colleague do with this brief?' People answer that more honestly than they answer about themselves, "
        "and the gap between the two answers reveals the real norm.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Reveals whether AI skepticism is quality-based (fixable with better outputs) or identity-based "
        "(fixable with culture change). Pairing the self-report with a reference group question gives you "
        "both the stated attitude and a more truthful read on actual behavior.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 05 ─────────────────────────────────────────────────────────────
    method_header(doc, 5, "Semantic Differential Scale")
    divider(doc)
    para(doc,
         "A semantic differential presents bipolar adjective pairs at opposite ends of a scale. "
         "It captures the evaluative, potency, and activity dimensions of attitudes more richly than "
         "agree/disagree items — especially useful for brand perception, product experience, and emotional response research.",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTION")
    para(doc,
         "Q5.  Using the scales below, indicate where your experience with the AI platform falls "
         "between each pair of words. (Mark the circle that best reflects your perception.)",
         size=10.5, space_after=4)

    pairs = [("Confusing", "Intuitive"), ("Unreliable", "Dependable"), ("Threatening", "Empowering"),
             ("Rigid", "Flexible"), ("Impersonal", "Human"), ("Slow", "Fast")]
    sd = doc.add_table(rows=len(pairs)+1, cols=9)
    sd.style = "Table Grid"
    set_cell_bg(sd.cell(0, 0), TABLE_HEADER_HEX)
    sd.cell(0, 0).paragraphs[0].add_run("").font.size = Pt(9)
    for j in range(1, 8):
        set_cell_bg(sd.cell(0, j), TABLE_HEADER_HEX)
        r = sd.cell(0, j).paragraphs[0].add_run(str(j))
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        sd.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(sd.cell(0, 8), TABLE_HEADER_HEX)
    sd.cell(0, 8).paragraphs[0].add_run("").font.size = Pt(9)
    for i, (left, right) in enumerate(pairs):
        bg = TABLE_ROW1_HEX if i % 2 == 0 else TABLE_ROW2_HEX
        set_cell_bg(sd.cell(i+1, 0), bg)
        lr = sd.cell(i+1, 0).paragraphs[0].add_run(left)
        lr.bold = True; lr.font.size = Pt(9)
        sd.cell(i+1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for j in range(1, 8):
            set_cell_bg(sd.cell(i+1, j), bg)
            p = sd.cell(i+1, j).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("◯").font.size = Pt(10)
        set_cell_bg(sd.cell(i+1, 8), bg)
        rr = sd.cell(i+1, 8).paragraphs[0].add_run(right)
        rr.bold = True; rr.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    callout_table(doc, "WHY THIS METHOD",
        "Standard Likert scales are one-directional. A semantic differential captures the full perceptual "
        "space — a tool can be rated as fast AND impersonal, which is a very different profile than slow "
        "AND human. This matters for product positioning and UX communication strategy.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Ideal for brand perception, product experience, and emotional tone research. The profiles cluster "
        "cleanly in analysis — you can visualize the 'personality' of a product or service and compare it "
        "across user segments, time points, or against a competitor.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 06 ─────────────────────────────────────────────────────────────
    method_header(doc, 6, "Social Circle / Network Influence Item")
    divider(doc)
    para(doc,
         "Social circle items ask respondents about the behaviors or attitudes of people in their immediate "
         "network. Rooted in descriptive norm research and peer influence theory, these items capture the "
         "social context of behavior — which is often a stronger predictor of individual adoption than attitudes alone.",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTIONS")
    para(doc,
         "Q6a.  Think about the 5 colleagues you work with most closely. To the best of your knowledge, "
         "how many of them regularly use AI tools in their work?",
         size=10.5, space_after=3)
    option_list(doc, ["0 of 5", "1 of 5", "2 of 5", "3 of 5", "4 of 5", "5 of 5"])
    para(doc,
         "Q6b.  In the past month, has a colleague recommended an AI tool to you, or have you recommended one to a colleague?",
         size=10.5, space_after=3)
    option_list(doc, ["Yes — a colleague recommended one to me", "Yes — I recommended one to a colleague", "Both", "Neither"])
    para(doc,
         "Q6c.  How important is it to you that your colleagues view you as someone who stays current with new professional tools?",
         size=10.5, space_after=3)
    option_list(doc, ["Not at all important", "Slightly important", "Moderately important", "Very important"])

    callout_table(doc, "WHY THIS METHOD",
        "Adoption behavior is shaped by what people believe others are doing. Q6a measures perceived peer "
        "adoption, which predicts individual adoption better than most attitude items. The key analytical "
        "move is comparing self-report to reference group report: ask 'would you do X?' and then 'would "
        "your colleague do X?' People answer the second question more honestly. The gap between those two "
        "answers is where the real norm lives — and it tells you whether adoption resistance is personal "
        "or social.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "If peer adoption is the primary predictor, the intervention is a champion network, not more "
        "training. Social circle items let you make that call with data instead of assumption.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # ── Method 07 ─────────────────────────────────────────────────────────────
    method_header(doc, 7, "Open-Ended Item with Structured Probe")
    divider(doc)
    para(doc,
         "An open-ended item followed by a structured probe captures qualitative depth within a quantitative "
         "survey. The first item invites free expression; the probe anchors the response to a specific dimension. "
         "This is more efficient than a full qualitative interview for large samples while providing richer data than a closed item alone.",
         size=10, color=GRAY, space_after=6)

    q_label(doc, "EXAMPLE QUESTIONS")
    para(doc,
         "Q7a.  In your own words, what is the biggest barrier to your organization fully embracing AI tools "
         "in research and analysis work?  [Open text — 500 character limit]",
         size=10.5, space_after=4)

    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    set_cell_bg(box.cell(0, 0), "F8F8F8")
    box.cell(0, 0).paragraphs[0].add_run("\n\n\n").font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    para(doc, "Q7b.  [PROBE]  You mentioned a barrier. Which category does it best fit?", size=10.5, space_after=3)
    option_list(doc, [
        "Technical limitations (quality, accuracy, speed)",
        "Organizational culture or leadership support",
        "Training or skill gaps",
        "Data privacy or security concerns",
        "Unclear use cases for my type of work",
        "Other / cannot categorize",
    ])
    para(doc,
         "Q7c.  How long do you think it would take for this barrier to be meaningfully reduced at your organization?",
         size=10.5, space_after=3)
    option_list(doc, ["Less than 6 months", "6–12 months", "1–2 years", "More than 2 years", "It won't be"])

    callout_table(doc, "WHY THIS METHOD",
        "Pure open-ended items are hard to analyze at scale. The probe converts the qualitative response into "
        "a categorized data point without forcing the respondent into a box prematurely. Q7c adds a "
        "time-perception dimension that closed barrier items almost never capture but that matters enormously for planning.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "You get the richness of open text for discovery and the analyzability of closed categories for "
        "reporting. The time-horizon probe tells you how urgent respondents think the problem is — which sets "
        "realistic expectations for change management timelines.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    divider(doc)
    para(doc,
         "About this document: All examples are drawn from or inspired by professional research engagements "
         "at RAND Corporation, the Gates Foundation EMO study, and behavioral science consulting work at "
         "JP Morgan Private Bank and TD Bank. Items have been anonymized for portfolio use.",
         size=9, color=GRAY, space_after=4)
    para(doc, footer_text(), size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = BASE / "01_survey_research" / "Survey Methods Showcase.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT 2 — Interview Methods Showcase
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def build_interview_doc():
    doc = Document()
    set_doc_margins(doc)

    para(doc, "QUALITATIVE RESEARCH METHODS SHOWCASE", bold=True, size=22, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Six interview techniques — with example questions, probe sequences, rationale, and client value.",
         size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Jeremy Barajas  |  Behavioral Scientist & Mixed-Methods Researcher",
         bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "RAND Corporation  ·  University of Pennsylvania (M.B.D.S.)  ·  USC (B.A. Psychology, Honors)",
         size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    divider(doc)
    para(doc,
         "Six qualitative interview techniques inspired by my work at RAND Corporation, JP Morgan Private Bank, "
         "and TD Bank. Part I is a complete annotated interview guide. Methods 01–06 demonstrate individual "
         "question types with design rationale and client value.",
         size=10, space_after=10)

    q_label(doc, "CONTENTS")
    for item in [
        "Part I   —  Annotated Semi-Structured Interview Guide (full structure)",
        "Method 01  —  Grand Tour Question",
        "Method 02  —  Laddering / Why Probe",
        "Method 03  —  Critical Incident Technique",
        "Method 04  —  Hypothetical / Future Scenario",
        "Method 05  —  Think-Aloud / Cognitive Walkthrough",
        "Method 06  —  Member Check Question",
    ]:
        para(doc, item, size=10, space_after=2)

    doc.add_page_break()

    # Part I — Annotated Guide
    heading(doc, "PART I — Annotated Semi-Structured Interview Guide", level=1)
    divider(doc)
    para(doc,
         "A semi-structured interview uses a prepared guide as a scaffold — not a script. "
         "The interviewer follows the respondent's lead while ensuring key topics are covered. "
         "The annotations below explain the purpose of each section and what to listen for.",
         size=10, color=GRAY, space_after=6)

    guide_sections = [
        ("GUIDE METADATA", [
            ("Study", "AI Platform Adoption — Employee Listening Study"),
            ("Duration", "45–60 minutes"),
            ("Interviewer", "Jeremy Barajas, RAND Corporation"),
            ("Note", "This guide is a scaffold, not a script. Follow the respondent's lead. "
                     "Probe on what's interesting or unexpected. You do not need to ask every question."),
        ]),
        ("SECTION 1 — Introduction & Consent  (3–5 min)", [
            ("Opening", "Thank you for taking the time to speak with me today. My name is Jeremy and I'm "
                        "a researcher studying how people at this organization experience AI tools in their work."),
            ("Ground rules", "There are no right or wrong answers. I'm here to learn from your perspective — "
                             "your honest experience is exactly what we need."),
            ("Recording", "With your permission, I'd like to record our conversation so I can focus on listening "
                          "rather than note-taking. The recording will only be used for my notes."),
            ("Interviewer note", "Wait for verbal consent before recording. Allow silence after consent — "
                                 "let the respondent settle in before the first question."),
        ]),
        ("SECTION 2 — Warm-Up  (5 min)", [
            ("Q", "Can you tell me a little about your role here — what you work on day-to-day?"),
            ("Purpose", "Establishes rapport and gives the interviewer context for interpreting later responses. "
                        "Listen for task types that AI might be relevant to."),
            ("Q", "How long have you been in this role, and how has your work changed over that time?"),
            ("Purpose", "Opens a historical frame that makes it natural to discuss change — including new tools."),
        ]),
        ("SECTION 3 — Core Questions  (25–30 min)", [
            ("Grand Tour", "Walk me through a typical week of analytical work. What does that look like from start to finish?"),
            ("Listen for", "Where time is spent, what feels tedious vs. meaningful, any tools mentioned. Do not interrupt."),
            ("Probe", "You mentioned [X task]. Tell me more about how you do that."),
            ("Critical Incident", "Think of a specific time recently when you had a lot of data to process or a tight deadline. What happened?"),
            ("Listen for", "Coping strategies, workarounds, frustration points — these are the highest-signal moments."),
            ("Probe", "What did you wish you had in that moment that you didn't have?"),
            ("Transition to AI", "Have you used any AI tools in your work — things like ChatGPT, Copilot, or anything similar?"),
            ("Branch A (used AI)", "Tell me about the first time you used one of those tools for work. What were you trying to do?"),
            ("Branch B (not used)", "What's kept you from trying one of those tools? [Do NOT prompt with a list of barriers.]"),
        ]),
        ("SECTION 4 — Deep Dive  (10–15 min)", [
            ("Laddering", "You said [reason for/against AI use]. Why does that matter to you? [Repeat 3x to reach value level.]"),
            ("Hypothetical", "Imagine the tool was perfect — it never made a mistake and you could trust it completely. "
                             "Would you use it more? For what?"),
            ("Vignette", "[Show scenario card] A colleague sends you a 10-page brief drafted by AI and lightly edited. "
                         "Would you use it? What would you do with it?"),
            ("Listen for", "Trust differentiation (output quality vs. source), social/reputational concern, calibration strategies."),
        ]),
        ("SECTION 5 — Close & Member Check  (5 min)", [
            ("Member Check", "Based on what you've shared, it sounds like [your summary]. Does that capture it accurately?"),
            ("Interviewer note", "This is the most important question. It surfaces misinterpretations before they become findings."),
            ("Final Q", "Is there anything else about your experience with AI tools that we haven't covered?"),
            ("Theoretical sampling", "Is there someone else at this organization you think I should talk to?"),
            ("Note", "Ask who would see things differently, not just who agrees."),
        ]),
    ]

    section_colors = [
        "555555", TABLE_HEADER_HEX, "2E7D52", "2D6A9F", "6A4C93", "555555"
    ]
    for (section_title, rows), bg in zip(guide_sections, section_colors):
        tbl = doc.add_table(rows=1+len(rows), cols=2)
        tbl.style = "Table Grid"
        set_cell_bg(tbl.cell(0, 0), bg)
        set_cell_bg(tbl.cell(0, 1), bg)
        merged = tbl.cell(0, 0).merge(tbl.cell(0, 1))
        hr_run = merged.paragraphs[0].add_run(section_title)
        hr_run.bold = True; hr_run.font.size = Pt(10); hr_run.font.color.rgb = WHITE
        for i, (label, value) in enumerate(rows):
            row_bg = TABLE_ROW1_HEX if i % 2 == 0 else TABLE_ROW2_HEX
            c0, c1 = tbl.cell(i+1, 0), tbl.cell(i+1, 1)
            set_cell_bg(c0, row_bg)
            set_cell_bg(c1, row_bg)
            c0.width = Inches(1.3)
            lr = c0.paragraphs[0].add_run(label)
            lr.bold = True; lr.font.size = Pt(9); lr.font.color.rgb = DARK
            vr = c1.paragraphs[0].add_run(value)
            vr.font.size = Pt(9); vr.font.color.rgb = DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # Method 01 — Grand Tour
    method_header(doc, 1, "Grand Tour Question")
    divider(doc)
    para(doc,
         "A grand tour question invites the respondent to narrate their world without the interviewer "
         "imposing categories. It is always the first substantive question. The topics, vocabulary, and "
         "sequence the respondent chooses are themselves data.",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "EXAMPLE QUESTION")
    para(doc,
         '"Walk me through a typical week of analytical work — from the moment you start a project to '
         'when you hand something off. What does that actually look like?"',
         size=10.5, italic=True, space_after=4)
    q_label(doc, "FOLLOW-ON PROBES")
    probe_list(doc, [
        "You mentioned [X] — tell me more about how you do that.",
        "What part of that takes the most time?",
        "Is there anything in that process that feels like a workaround?",
        "What would an ideal version of that process look like?",
    ])
    callout_table(doc, "WHY THIS METHOD",
        "Imposes no categories on the respondent. You find out what matters to them, not just whether "
        "what matters to you matters to them. The grand tour reveals the workflow structure, vocabulary, "
        "and priorities that inform all subsequent probing.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Surfaces needs, frustrations, and workarounds the client didn't know to ask about. "
        "The most valuable insights in needfinding almost always come from grand tour responses, "
        "not from pre-specified question topics.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Method 02 — Laddering
    method_header(doc, 2, "Laddering / Why Probe")
    divider(doc)
    para(doc,
         "Laddering moves from a surface-level belief or behavior to the underlying value driving it "
         "by repeatedly asking 'why does that matter to you?' Each iteration climbs one rung from "
         "attribute → consequence → terminal value. Three iterations usually reaches the value level.",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "EXAMPLE SEQUENCE")
    spec_table(doc, [
        ("Turn", "Question / Expected Response Level"),
        ("1", "\"You said you don't fully trust the AI outputs. Why does that matter to you?\"  →  Attribute: 'it makes mistakes'"),
        ("2", "\"And why does it matter that it makes mistakes in your context?\"  →  Consequence: 'I'd have to check everything'"),
        ("3", "\"Why is having to check everything a problem for you?\"  →  Consequence: 'it defeats the purpose'"),
        ("4", "\"What's the purpose you'd want it to serve?\"  →  Terminal value: 'I want to trust my own work'"),
    ], col_widths=(0.6, 5.9))
    callout_table(doc, "WHY THIS METHOD",
        "Surface reasons are rarely actionable. Knowing someone dislikes AI 'because it makes mistakes' "
        "tells you nothing about what to do. Knowing they distrust it because it threatens their sense of "
        "professional identity tells you exactly what the intervention needs to address.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Reveals the real reason for resistance — which is almost never the stated reason. Interventions "
        "designed against surface reasons fail. Laddering gives you the level at which change actually needs to happen.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # Method 03 — Critical Incident
    method_header(doc, 3, "Critical Incident Technique")
    divider(doc)
    para(doc,
         "The Critical Incident Technique (CIT) asks respondents to recount a specific, memorable event "
         "rather than summarizing their general experience. People are far more accurate about concrete events "
         "than about abstract impressions. The episode structure (situation → action → outcome) generates rich, analyzable data.",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "EXAMPLE QUESTIONS")
    para(doc, '"Think of a specific time recently when you used an AI tool and it went really well. '
              'Walk me through exactly what happened."', size=10.5, italic=True, space_after=4)
    para(doc, '"Now think of a time it didn\'t go well, or when you tried to use it and gave up. '
              'What was happening? What did you try? What made you stop?"', size=10.5, italic=True, space_after=4)
    q_label(doc, "KEY PROBES")
    probe_list(doc, [
        "What exactly did you do next?  (keeps the narrative specific)",
        "What were you thinking at that moment?  (surfaces the cognitive layer)",
        "What would you have done differently?  (reveals latent preferences)",
        "Has that happened more than once?  (tests whether this is typical or anomalous)",
    ])
    callout_table(doc, "WHY THIS METHOD",
        "General questions produce general answers. A specific incident produces a behavior sequence you "
        "can map, code, and compare across participants. CIT is especially powerful for identifying failure "
        "modes and unmet needs that respondents have normalized and would never surface unprompted.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Generates concrete use cases, failure scenarios, and success conditions that product teams, "
        "trainers, and communicators can act on directly. Much higher signal than 'what do you think about X?' questions.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Method 04 — Hypothetical
    method_header(doc, 4, "Hypothetical / Future Scenario")
    divider(doc)
    para(doc,
         "A hypothetical removes a known constraint to reveal whether that constraint is the real barrier "
         "or whether reluctance lives deeper. It is also used to elicit desired future states: what would "
         "the ideal look like if the current limitation didn't exist?",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "EXAMPLE QUESTIONS")
    for q in [
        "\"Imagine the AI tool was perfect — it never made a factual error and you could trust every output completely. How would that change how you use it, if at all?\"",
        "\"If you could redesign your analytical workflow from scratch — no legacy systems, no organizational constraints — what would it look like?\"",
        "\"Five years from now, what do you think your job looks like if AI tools have become standard in your field?\"",
    ]:
        para(doc, q, size=10.5, italic=True, space_after=3)
    callout_table(doc, "WHY THIS METHOD",
        "Hypotheticals are a diagnostic tool. If removing the stated barrier doesn't change the answer — "
        "'I still wouldn't trust it even if it were perfect' — then the real issue is identity, culture, or "
        "relevance, not quality. The hypothetical cleanly separates symptom from root cause.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Reveals unspoken objections and latent demand. 'I'd use it for X but not Y' from a hypothetical "
        "question is high-signal product feedback that no closed survey item would produce.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    doc.add_page_break()

    # Method 05 — Think-Aloud
    method_header(doc, 5, "Think-Aloud / Cognitive Walkthrough")
    divider(doc)
    para(doc,
         "In a think-aloud, the respondent performs a real task while narrating their thought process. "
         "The interviewer observes and asks minimal, non-leading questions. It is the gold standard for "
         "understanding how someone actually uses a tool or interface — not how they remember using it.",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "SETUP")
    para(doc,
         '"I\'m going to ask you to use the AI platform to complete a task you\'d normally do in your work. '
         'As you do it, please talk out loud — tell me what you\'re thinking, what you\'re looking at, and '
         'what you expect to happen. There\'s no right way to do this. I\'m not testing you — I\'m testing the tool."',
         size=10.5, italic=True, space_after=6)
    q_label(doc, "DURING-TASK PROBES  (minimal intervention)")
    probe_list(doc, [
        "What are you looking for right now?",
        "What did you expect to happen there?",
        "You paused — what were you thinking?",
        "What would you normally do at this point?  (if they get stuck)",
    ])
    q_label(doc, "POST-TASK PROBES")
    probe_list(doc, [
        "Was there a moment where you felt confused or uncertain?",
        "What would have made that easier?",
        "Is that how you'd normally approach this task, or did you do it differently because I was watching?",
    ])
    callout_table(doc, "WHY THIS METHOD",
        "People cannot accurately self-report their cognitive processes after the fact. Think-aloud captures "
        "the real-time confusion, hesitation, and workaround behavior that post-hoc interviews miss. "
        "It is the gold standard for UX research and cognitive task analysis.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Surfaces specific UI/UX failure points, mental model mismatches, and comprehension gaps with "
        "precision a survey cannot achieve. Every pause is a data point. Used in 16 cognitive interviews "
        "for JP Morgan Private Bank and 9 sessions for TD Bank — both produced actionable product recommendations.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Method 06 — Member Check
    method_header(doc, 6, "Member Check Question")
    divider(doc)
    para(doc,
         "A member check asks the respondent to verify the interviewer's interpretation of what they've said. "
         "It catches misinterpretations before they become findings, and it often produces the most honest, "
         "corrective language in the entire interview.",
         size=10, color=GRAY, space_after=6)
    q_label(doc, "EXAMPLE QUESTIONS")
    for q in [
        "\"Based on what you've shared, it sounds like the core concern is accountability — being able to defend the work if someone questions it. Does that capture it?\"",
        "\"When you say you 'don't trust it' — do you mean the outputs are wrong, or that you can't tell whether they're wrong?\"",
        "\"If I were to summarize your experience as: [summary], would you say that's accurate? What would you change about that?\"",
    ]:
        para(doc, q, size=10.5, italic=True, space_after=3)
    callout_table(doc, "WHY THIS METHOD",
        "Researcher interpretation bias is the most common source of error in qualitative research. "
        "The member check is the most direct check on it. A respondent who says 'no, that's not quite right' "
        "has just given you your finding. Without the check, you'd have published the wrong conclusion.",
        bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
    callout_table(doc, "CLIENT VALUE",
        "Increases credibility of findings. When you can tell a client 'we ran this interpretation back "
        "to 28 participants and they confirmed it,' the recommendation carries significantly more weight "
        "than findings derived entirely from the researcher's reading of transcripts.",
        bg_hex=ORANGE_LIGHT_HEX, head_color=ORANGE)

    divider(doc)
    para(doc,
         "About this document: Techniques draw from professional research at RAND Corporation, "
         "the Gates Foundation EMO study (cognitive interviews), JP Morgan Private Bank (16 cognitive "
         "interviews), and TD Bank (9 UX needfinding sessions). All content anonymized for portfolio use.",
         size=9, color=GRAY, space_after=4)
    para(doc, footer_text(), size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = BASE / "05_qualitative_methods" / "Interview Methods Showcase.docx"
    doc.save(out)
    print(f"Saved: {out}")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# DOCUMENT 3 — Experiment Design Showcase
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def build_experiment_doc():
    doc = Document()
    set_doc_margins(doc)

    para(doc, "EXPERIMENTAL & RESEARCH DESIGN SHOWCASE", bold=True, size=22, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "Five research design approaches — with structure, hypotheses, power analysis, and validity threats.",
         size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "Jeremy Barajas  |  Behavioral Scientist & Mixed-Methods Researcher",
         bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "RAND Corporation  ·  University of Pennsylvania (M.B.D.S.)  ·  USC (B.A. Psychology, Honors)",
         size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    divider(doc)
    para(doc,
         "Five experimental and quasi-experimental research designs, each built around a realistic "
         "business or consulting problem. Includes design logic, worked example, power analysis, "
         "and validity threats.",
         size=10, space_after=10)

    q_label(doc, "CONTENTS")
    for item in [
        "Design 01  —  Between-Subjects Randomized Controlled Trial (RCT)",
        "Design 02  —  Within-Subjects / Repeated Measures",
        "Design 03  —  Survey Experiment (Vignette Factorial Design)",
        "Design 04  —  Pre/Post with Control Group (Quasi-Experimental, DiD)",
        "Design 05  —  Behavioral Nudge A/B Test (Three-Arm)",
    ]:
        para(doc, item, size=10, space_after=2)

    doc.add_page_break()

    designs = [
        (1, "Between-Subjects Randomized Controlled Trial (RCT)", [
            "The RCT is the gold standard for causal inference. Each participant sees only one condition, "
            "and random assignment eliminates systematic confounding between groups.",
            [
                ("Business problem", "An e-commerce client wants to know: does showing customer reviews on the checkout page reduce cart abandonment, or does it distract buyers and increase drop-off?"),
                ("Research question", "Does displaying social proof (star ratings + review count) on the checkout page increase purchase completion relative to no social proof?"),
                ("Hypothesis", "Users in the social proof condition will complete purchase at a higher rate than users in the control condition. (One-tailed, α=.05.)"),
                ("Conditions", "Control: checkout page with no reviews (n=500)\nTreatment: checkout page with star rating + review count displayed (n=500)"),
                ("Randomization", "Cookie-level randomization; users held in their assigned condition across sessions."),
                ("Primary outcome", "Purchase completion within the session (binary)."),
                ("Analysis", "Chi-square test on completion rate; logistic regression controlling for device type, traffic source, and cart value."),
                ("Power", "Detecting Δ=3pp (12%→15% completion) with 80% power at α=.05 requires n=430/group. N=500/group provides attrition buffer."),
            ],
            "Only an RCT establishes that social proof caused the difference. Without randomization, "
            "high-intent buyers might self-select into seeing reviews through other pathways, "
            "confounding the estimate.",
            "Requires sufficient traffic volume to reach target N within a reasonable test window. "
            "Seasonal effects can bias results if treatment and control periods do not overlap.",
        ]),
        (2, "Within-Subjects / Repeated Measures", [
            "Each participant experiences all conditions, eliminating between-person variance and "
            "dramatically increasing statistical power. Requires careful sequencing to prevent "
            "earlier conditions from contaminating later ones.",
            [
                ("Business problem", "A SaaS company suspects users rate their product differently depending on whether they see a competitor's interface first. They want to measure the framing effect without running two separate studies."),
                ("Research question", "Does exposure order (own product first vs. competitor first) affect perceived usability ratings?"),
                ("Hypothesis", "Users who see the competitor interface first will rate the client's product higher than users who see it first, due to contrast effects."),
                ("Conditions", "Each participant rates both interfaces. Order counterbalanced: half see client first, half see competitor first."),
                ("Stimuli", "Identical task scenarios administered in both interfaces. 10-minute washout task between sessions."),
                ("Outcome", "Usability rating (1–7 SUS-adapted scale); task completion time."),
                ("Analysis", "Repeated-measures ANOVA; order as between-subjects factor to test for contrast effect."),
                ("Power", "Within-subjects correlation ≈.5; Cohen's d=0.4 detectable with n=52 at 80% power."),
            ],
            "Each participant serves as their own control. Individual differences in general "
            "usability standards drop out of the analysis, leaving a cleaner estimate of the order effect.",
            "Demand characteristics: participants may adjust ratings after seeing both products. "
            "Not appropriate when product exposure has a permanent learning effect.",
        ]),
        (3, "Survey Experiment (Vignette Factorial Design)", [
            "A survey experiment embeds a randomized manipulation inside a survey. Respondents are "
            "randomly assigned to different versions of a scenario. Useful when the behavior of interest "
            "is sensitive, rare, or difficult to observe directly.",
            [
                ("Business problem", "A restaurant chain wants to know whether calorie labeling increases healthy menu selection — but they can't run a field experiment across locations without tipping off staff and changing service behavior."),
                ("Research question", "Does displaying calorie counts on a menu (vs. not) affect stated meal selection, and does the effect differ by customer health motivation?"),
                ("Design", "2×2 factorial. Factor A: Calorie label (present vs. absent). Factor B: Customer type (health-motivated vs. not). N=150 per cell, N=600 total."),
                ("Main effects", "H₁: Calorie labels shift selection toward lower-calorie items.\nH₂: Health-motivated customers show stronger label effect."),
                ("Interaction", "H₃: Label × motivation interaction — labels may only matter for already health-motivated customers."),
                ("Outcome", "Stated meal choice (calories); willingness to pay for healthy option."),
                ("Analysis", "2x2 ANOVA; effect sizes (η²); planned contrasts for interaction."),
                ("Power", "Main effect d=0.25 detectable with 80% power at n=100/cell. 150/cell provides power for interaction."),
            ],
            "Separates the label effect from customer motivation in a single study. Field data can't "
            "disentangle these because health-motivated customers already choose differently.",
            "Stated choices may overestimate real behavior (social desirability). Vignettes simplify "
            "real ordering contexts. Interaction effects require larger N than main effects.",
        ]),
        (4, "Pre/Post with Control Group (Quasi-Experimental, DiD)", [
            "When randomization isn't feasible, the difference-in-differences (DiD) framework uses a "
            "non-equivalent control group to remove shared time trends from the causal estimate.",
            [
                ("Business problem", "A retail chain rolled out a new employee incentive program in 12 stores but not in 8 others due to budget. Leadership wants to know if the program increased sales — but didn't randomize which stores got it."),
                ("Research question", "Did the incentive program increase monthly revenue in treated stores, net of trends affecting all stores?"),
                ("DiD estimate", "Treated stores: +$18K/month avg (pre→post). Control stores: +$6K/month avg (same period). DiD causal estimate: +$12K/month attributable to the program."),
                ("Key assumption", "Parallel trends: treated and control stores would have grown at the same rate without the program. Validated by plotting the 6 months prior to rollout."),
                ("Threats", "Selection bias (high-performing stores may have been chosen for the program); spillover if treated-store staff share tactics with control-store staff."),
                ("Analysis", "DiD regression with store and month fixed effects; cluster-robust standard errors at store level."),
                ("Data required", "Monthly revenue for all stores for 6+ months before and after rollout; program assignment records."),
            ],
            "Extracts a credible causal estimate from a rollout that wasn't randomized. The control "
            "stores show what would have happened to treated stores without the program.",
            "Parallel trends assumption is untestable in the post-period. Results are sensitive to "
            "which stores were selected for treatment and why.",
        ]),
        (5, "Behavioral Nudge A/B Test (Three-Arm)", [
            "A nudge A/B test applies choice architecture — social proof, implementation intentions, "
            "loss framing — to shift behavior at scale. A multi-arm design identifies which mechanism "
            "is working, not just whether any message helps.",
            [
                ("Business problem", "A subscription app has a large free-tier user base but low conversion to paid. Email campaigns have had minimal effect. The client wants to know which message framing drives the most upgrades."),
                ("Research question", "Which email framing most increases free-to-paid conversion within 7 days?"),
                ("Arm A (Control)", "Standard: 'Upgrade to unlock all features.'"),
                ("Arm B — Social proof", "'Join the 62% of users who upgraded last month.'  (Descriptive norm — Cialdini, 1984)"),
                ("Arm C — Impl. intention", "'Pick a time this week to try premium free for 7 days. [Choose a time]'  (Planning prompt — Gollwitzer, 1999)"),
                ("Arm D — Loss frame", "'Your free access to [Feature X] expires in 3 days.'  (Loss aversion — Kahneman & Tversky, 1979)"),
                ("Arm E — Fresh Start", "'The new month starts in 3 days — a great time to try premium free.'  (Temporal landmark — Milkman & Dai, 2014)"),
                ("Primary outcome", "Free-to-paid conversion within 7 days of email (binary)."),
                ("Analysis", "Chi-square omnibus test across arms; pairwise comparisons with Bonferroni correction (α=.05/4=.0125 per arm vs. control)."),
                ("Power", "Baseline conversion ~20%. Detecting Δ=10pp with 80% power requires n=176/arm. N=200/arm across 5 arms provides buffer."),
                ("Ethics note", "Loss framing arm requires legal/HR review — language implying penalty must be accurate and pre-approved."),
            ],
            "The five-arm design maps each treatment to a distinct behavioral mechanism — descriptive norms "
            "(Cialdini), planning prompts (Gollwitzer), loss aversion (Kahneman & Tversky), and temporal "
            "landmarks (Milkman & Dai). Testing them head-to-head identifies which lever actually drives "
            "conversion for this population, not just whether any message helps.",
            "Nudge effects often attenuate over time (habituation). This test measures 7-day conversion — "
            "sustained behavior change requires follow-up measurement. Social proof messages require "
            "accurate figures; fabricated norms backfire. Fresh Start effects depend on the user perceiving "
            "a genuine temporal landmark, which varies by individual.",
        ]),
    ]

    for num, title, content in designs:
        design_header(doc, num, title)
        divider(doc)
        para(doc, content[0], size=10, color=GRAY, space_after=6)
        q_label(doc, "WORKED EXAMPLE / DESIGN SPECIFICATION")
        spec_table(doc, content[1])
        callout_table(doc, "WHY THIS DESIGN", content[2], bg_hex=BLUE_LIGHT_HEX, head_color=BLUE)
        callout_table(doc, "LIMITATIONS", content[3], bg_hex="FAEAEC", head_color=RGBColor(0x9B, 0x23, 0x35))
        if num < 5:
            doc.add_page_break()

    divider(doc)
    para(doc,
         "About this document: Design examples are informed by published research and professional evaluation "
         "work at RAND Corporation, including the RAND-EDA AI adoption study, FEMA building code adoption "
         "research, and the Gates Foundation behavioral segmentation project. All examples are illustrative. "
         "Full experimental design, IRB coordination, power analysis, and statistical analysis available as project services.",
         size=9, color=GRAY, space_after=4)
    para(doc, footer_text(), size=8, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = BASE / "06_experiment_design" / "Experiment Design Showcase.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_survey_doc()
    build_interview_doc()
    build_experiment_doc()
    print("\nAll three Word documents built.")
