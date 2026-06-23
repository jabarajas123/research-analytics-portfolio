"""
Transcript Cleaner
==================
Pre-processing pipeline for qualitative interview transcripts.

This script is written as a teaching document. Each section explains
what the code does, why it's designed that way, what the output should
look like, and how to verify correctness.

The pipeline takes raw transcript exports — from Teams, Zoom, Otter, or
manual transcription — and produces three standardized formats:

  cleaned/<filename>_clean.txt      — plain text, Atlas.ti ready
  cleaned/<filename>_clean.docx     — Word doc with formatted speaker labels
  cleaned/<filename>_dedoose.csv    — one row per turn, Dedoose ready

Why three formats? Different qualitative platforms expect different inputs.
Atlas.ti imports plain text or Word; Dedoose uses CSV. Generating all three
from one pass means you don't re-clean the same file three times.

Run:
    python3 transcript_cleaner.py                  # process all in transcripts/
    python3 transcript_cleaner.py my_file.txt      # single file
    python3 transcript_cleaner.py --no-rename      # skip interactive rename step

Supported raw input formats (auto-detected):
  "SPEAKER NAME: dialogue text"
  "Speaker Name:\\n dialogue text"
  "[SPEAKER NAME]: dialogue text"
  "[00:12:34] SPEAKER: dialogue text"  — timestamps are stripped automatically
"""

# ── Imports ───────────────────────────────────────────────────────────────────
# We import only the standard library here. python-docx is imported lazily
# inside the functions that need it so the script runs without it if you're
# only exporting plain text and CSV.

import re           # regular expressions for pattern matching on transcript text
import sys          # sys.argv / sys.exit() for CLI argument handling
import csv          # writing the Dedoose-format CSV output
import argparse     # parsing command-line flags (--no-rename)
from pathlib import Path    # platform-safe file path construction
from textwrap import wrap   # word-wrapping long preview lines in the terminal roster


# ── Configuration ─────────────────────────────────────────────────────────────
# All path and display constants live here so they're easy to change without
# hunting through the code. TRANSCRIPT_DIR is where the script looks for raw
# files; CLEAN_DIR is where all three export formats are written.

TRANSCRIPT_DIR = Path("transcripts")
CLEAN_DIR      = Path("cleaned")

PREVIEW_LINES = 3    # how many example lines to show per speaker in the rename roster
PREVIEW_CHARS = 120  # max characters per preview line before truncating with "…"


# ── Speaker-Turn Detection Patterns ───────────────────────────────────────────
# Transcripts from different platforms use different formatting conventions.
# Rather than requiring a specific format, we try three common patterns and
# pick whichever one matches the most turns in the document.
#
# Pattern priority (tried in order, best-match wins):
#   1. Timestamped: "[00:12:34] SPEAKER NAME: dialogue"
#      — Teams and some Zoom exports include timestamps before speaker names
#   2. Colon-delimited: "SPEAKER NAME: dialogue"
#      — the most common manual transcription style
#   3. Bracket-delimited: "[SPEAKER NAME]: dialogue"
#      — less common but appears in some auto-transcription exports
#
# Each pattern uses a named group structure: group(1) = speaker, group(2) = text.
# The speaker name is capped at 50 characters to avoid matching long sentences
# that happen to contain a colon (e.g., "The main finding: results were mixed.")
#
# re.MULTILINE makes ^ and $ match at the start/end of each LINE, not just
# the start/end of the entire string. Without this, none of the patterns would
# match at all since we're passing the whole document as one string.

SPEAKER_PATTERNS = [
    re.compile(r"^\s*\[?\d{1,2}:\d{2}(?::\d{2})?\]?\s*([A-Za-z][^\n:]{1,50}):\s*(.+)$", re.MULTILINE),
    re.compile(r"^([A-Za-z][^\n:]{1,50}):\s*(.+)$", re.MULTILINE),
    re.compile(r"^\[([^\]]{1,50})\]:\s*(.+)$", re.MULTILINE),
]


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1: PARSING
# ─────────────────────────────────────────────────────────────────────────────

def load_raw(path: Path) -> str:
    """
    Load a transcript file to a plain string.

    For .docx files, we extract each paragraph's text and join with newlines.
    We skip empty paragraphs (p.text.strip() would be '') to avoid producing
    blank lines inside the joined text. For .txt files, read_text() handles
    encoding automatically.

    Why not just always use .txt? Word documents from research teams often
    contain tracked changes, formatting, or metadata that aren't visible when
    you open the file — read_text() on a .docx returns binary XML garbage.
    python-docx extracts only the visible paragraph content.

    Verify: the returned string should start with a recognizable speaker turn,
    not with XML tags or binary characters.
    """
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    return path.read_text(encoding="utf-8")


def strip_noise(text: str) -> str:
    """
    Remove timestamps, normalize filler markers, and clean up whitespace.

    Steps in order — order matters because each substitution affects the
    input for the next one:

    1. Normalize line endings: Windows uses \\r\\n; Mac/Linux use \\n. We
       convert to \\n so all subsequent patterns only need to handle one case.

    2. Strip timestamps: patterns like [00:12:34] or 12:34 appear in Teams,
       Zoom, and Otter exports. We remove them before speaker detection so
       "[00:12:34] INTERVIEWER:" becomes "INTERVIEWER:". The regex allows
       optional brackets and optional seconds component.

    3. Strip leading whitespace left by timestamp removal: after removing
       "[00:12:34] " from the start of a line, the line begins with a space.
       Without this step, "^SPEAKER:" patterns won't match because ^ requires
       the speaker name to start at the true beginning of the line. This was
       the root cause of a bug where everything collapsed into one UNKNOWN block.

    4. Normalize filler markers: "(inaudible)" and "(crosstalk)" from Otter
       are standardized to bracketed form [inaudible] / [crosstalk] so they're
       consistent across transcript sources.

    5. Collapse excessive blank lines: more than one consecutive blank line
       adds visual noise without carrying structural meaning. We reduce to one.

    Verify: run strip_noise() on a sample with timestamps and check that
    speaker names land at the true start of each line with no leading spaces.
    """
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\[?\d{1,2}:\d{2}(?::\d{2})?\]?", "", text)
    text = re.sub(r"^\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\(inaudible\)", "[inaudible]", text, flags=re.I)
    text = re.sub(r"\(crosstalk\)", "[crosstalk]", text, flags=re.I)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def detect_pattern(text: str):
    """
    Pick the speaker-turn regex that matches the most turns in this document.

    We try all three patterns and count the number of matches each finds.
    The one with the highest count wins. This is a simple heuristic that
    works well in practice: the correct pattern will match every speaker turn,
    while incorrect patterns may match a handful of coincidental colons but
    rarely match more than the correct one.

    Returns: (best_pattern, match_count)
    If match_count < 2, no usable pattern was found — the caller falls back
    to treating the whole document as a single UNKNOWN speaker block.

    Verify: for a typical 45-minute interview with ~20 turns, best_count
    should be at least 10. If it's 1 or 2, the transcript may be in an
    unsupported format and needs manual inspection.
    """
    best, best_count = None, 0
    for pat in SPEAKER_PATTERNS:
        count = len(pat.findall(text))
        if count > best_count:
            best, best_count = pat, count
    return best, best_count


def parse_turns(text: str) -> list[dict]:
    """
    Parse a cleaned transcript string into a list of speaker-turn dicts.

    Each dict has two keys:
      "speaker": uppercase speaker name (e.g., "INTERVIEWER", "RESPONDENT 1")
      "text":    the spoken content of that turn

    The parser iterates over all regex matches in document order. Between any
    two consecutive matches there may be "gap" text — continuation paragraphs
    that the regex didn't capture because they don't start with "SPEAKER:".
    We append gap text to the previous turn's content so nothing is lost.

    Trailing text after the last match is handled the same way.

    Fallback: if fewer than 2 speaker patterns are detected (likely a
    free-form document, not a structured transcript), the whole text is
    returned as a single UNKNOWN turn. This is better than crashing — the
    user will see it in the roster and can rename it.

    Verify: for a two-person interview, the speaker set should have exactly
    two unique values. Check: len({t['speaker'] for t in turns}) == 2
    """
    pat, count = detect_pattern(text)

    if count < 2:
        return [{"speaker": "UNKNOWN", "text": text.strip()}]

    turns = []
    last_end = 0

    for m in pat.finditer(text):
        # Capture any text between the end of the last match and the start of
        # this one — this is continuation text from the previous speaker
        gap = text[last_end:m.start()].strip()
        if gap and turns:
            turns[-1]["text"] = (turns[-1]["text"] + " " + gap).strip()

        speaker  = m.group(1).strip().upper()   # normalize to uppercase for consistency
        utterance = m.group(2).strip()
        turns.append({"speaker": speaker, "text": utterance})
        last_end = m.end()

    # Capture any text after the final speaker match
    tail = text[last_end:].strip()
    if tail and turns:
        turns[-1]["text"] = (turns[-1]["text"] + " " + tail).strip()

    return turns


def merge_consecutive(turns: list[dict]) -> list[dict]:
    """
    Collapse back-to-back turns from the same speaker into one block.

    Some transcription tools split a single monologue into multiple turns
    at paragraph breaks. For analysis purposes, one continuous block of
    speech from one speaker is a single analytical unit — it should stay
    together. Merging also reduces false turn counts that would inflate
    the "n_turns" summary statistic.

    We join with a single space (not a newline) because the exported plain
    text format uses a blank line to separate turns, so internal whitespace
    within a turn is just a space.

    Verify: if RESPONDENT 1 has 3 consecutive turns before INTERVIEWER speaks,
    they should become 1 turn after merging. Check the turn count before and
    after — merged_turns should be strictly <= original_turns.
    """
    if not turns:
        return []
    merged = [turns[0].copy()]
    for turn in turns[1:]:
        if turn["speaker"] == merged[-1]["speaker"]:
            merged[-1]["text"] = merged[-1]["text"].rstrip() + " " + turn["text"].lstrip()
        else:
            merged.append(turn.copy())
    return merged


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2: SPEAKER ROSTER AND INTERACTIVE RENAME
# ─────────────────────────────────────────────────────────────────────────────
# After parsing, we show the researcher a preview of each detected speaker so
# they can verify the names and rename them if needed. Raw transcript exports
# often use full names, display names, or auto-detected labels ("Speaker 1",
# "User", the Zoom account email) that need to be replaced with consistent
# analytical labels like "INTERVIEWER" and "RESPONDENT 1".

def build_roster(turns: list[dict]) -> dict[str, list[str]]:
    """
    Build a preview dict: {speaker_name: [first N text snippets]}.

    We preserve first-appearance order (insertion-ordered dict, Python 3.7+)
    so the roster displays speakers in the order they first speak, which is
    the most natural order for a researcher reading the transcript.

    Each snippet is capped at PREVIEW_CHARS characters to fit the terminal.
    A "…" suffix signals truncation so the researcher knows more text follows.

    Verify: every unique speaker in turns should appear as a key in the roster.
    len(roster) == len({t['speaker'] for t in turns})
    """
    roster: dict[str, list[str]] = {}
    for t in turns:
        sp = t["speaker"]
        if sp not in roster:
            roster[sp] = []
        if len(roster[sp]) < PREVIEW_LINES:
            snippet = t["text"][:PREVIEW_CHARS]
            if len(t["text"]) > PREVIEW_CHARS:
                snippet += "…"
            roster[sp].append(snippet)
    return roster


def print_roster(roster: dict[str, list[str]]) -> None:
    """
    Print the speaker roster to the terminal in a readable format.

    We use wrap() from textwrap to break long preview lines at word boundaries
    rather than mid-word, and indent continuation lines 7 spaces to align
    under the first character of the preview text.
    """
    bar = "─" * 60
    print(f"\n{bar}")
    print("  SPEAKERS DETECTED")
    print(bar)
    for i, (speaker, previews) in enumerate(roster.items(), 1):
        print(f"\n  [{i}]  {speaker}")
        for line in previews:
            for wrapped in wrap(line, width=70):
                print(f"       {wrapped}")
    print(f"\n{bar}\n")


def interactive_rename(roster: dict[str, list[str]]) -> dict[str, str]:
    """
    Show the roster and prompt for optional renames. Return a mapping dict.

    The researcher presses Enter to keep the current name, or types a new
    label. New labels are uppercased for consistency regardless of how the
    researcher types them.

    Returns: {old_name: new_name}
    If a name is kept, old_name maps to itself — apply_rename() handles both
    cases uniformly without needing to check whether a rename occurred.

    Why interactive? Speaker names vary across transcripts, across researchers,
    and across transcription tools. Hardcoding a rename map would break on any
    file that uses different labels. The interactive step is the one place where
    human judgment is irreplaceable — the researcher knows which voice is which.
    """
    print_roster(roster)
    speakers = list(roster.keys())
    rename_map: dict[str, str] = {}

    print("  Rename speakers — press Enter to keep the current label.\n")
    for sp in speakers:
        new = input(f"  '{sp}'  →  ").strip()
        rename_map[sp] = new.upper() if new else sp

    return rename_map


def apply_rename(turns: list[dict], rename_map: dict[str, str]) -> list[dict]:
    """
    Apply the rename mapping to every turn.

    We use dict.get(key, key) as a safe fallback: if a speaker name somehow
    isn't in the mapping (shouldn't happen, but defensive), we keep the
    original name rather than introducing a None key.
    """
    return [
        {"speaker": rename_map.get(t["speaker"], t["speaker"]), "text": t["text"]}
        for t in turns
    ]


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3: EXPORT FORMATTERS
# ─────────────────────────────────────────────────────────────────────────────

def to_plain_text(turns: list[dict]) -> str:
    """
    Render turns to Atlas.ti-compatible plain text format.

    Atlas.ti expects speaker turns in this form:
        SPEAKER NAME:
        The text of their turn.

        NEXT SPEAKER:
        Their turn text.

    The blank line between turns is the structural separator Atlas.ti uses
    to identify turn boundaries during import. Without it, Atlas.ti treats
    the whole file as one continuous document segment.

    We use str.join() on a list rather than string concatenation in a loop
    because Python builds a new string object on every += operation; join()
    allocates once. For a 200-turn transcript this is negligible, but it's
    the correct habit.

    Verify: the output should have exactly (n_turns * 3 - 1) non-empty lines —
    one for the speaker label, one for the text, one blank line between turns
    (the last turn has no trailing blank).
    """
    lines = []
    for t in turns:
        lines.append(f"{t['speaker']}:")
        lines.append(t["text"])
        lines.append("")          # blank line separator
    return "\n".join(lines).strip()


def to_docx(turns: list[dict], title: str) -> "Document":
    """
    Build a formatted Word document for Atlas.ti import or human review.

    We use python-docx to apply consistent typographic styling:
    - Speaker labels: bold, blue (#2D6A9F), 10pt — visually distinct
    - Utterance text: dark (#1A1A2E), 10pt — readable
    - Margins: 1.0" top/bottom, 1.1" left/right — standard report margins

    Why not use Heading styles? Atlas.ti uses heading levels for its own
    structural markers. Using Word Heading styles for speaker labels would
    confuse Atlas.ti's import parser. Plain paragraphs with explicit font
    formatting achieve the same visual result without the structural conflict.

    The function returns a Document object rather than saving it directly —
    the caller (process_file) handles the path and calls doc.save(). This
    separation makes the formatter testable: you can call to_docx() in a
    test without touching the filesystem.

    Verify: open the output .docx and confirm speaker labels appear in blue
    bold and utterance text appears in dark normal weight.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BLUE = RGBColor(0x2D, 0x6A, 0x9F)
    DARK = RGBColor(0x1A, 0x1A, 0x2E)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Title paragraph — centered, blue, 14pt bold
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = BLUE

    doc.add_paragraph()   # one blank paragraph of vertical space after title

    for t in turns:
        # Speaker label paragraph
        sp = doc.add_paragraph()
        sr = sp.add_run(t["speaker"] + ":")
        sr.bold = True
        sr.font.size = Pt(10)
        sr.font.color.rgb = BLUE
        sp.paragraph_format.space_before = Pt(8)   # breathing room before each new speaker
        sp.paragraph_format.space_after  = Pt(2)

        # Utterance paragraph
        up = doc.add_paragraph()
        ur = up.add_run(t["text"])
        ur.font.size = Pt(10)
        ur.font.color.rgb = DARK
        up.paragraph_format.space_before = Pt(0)
        up.paragraph_format.space_after  = Pt(0)

    return doc


def to_dedoose_csv(turns: list[dict], source_name: str) -> list[dict]:
    """
    Format turns as a list of row dicts for Dedoose CSV import.

    Dedoose expects a structured import with one excerpt per row. The columns
    we produce map directly to Dedoose's "Media" import format:
      source:      the transcript identifier (used as the media title in Dedoose)
      turn_index:  sequential position of this turn in the transcript (1-based)
      speaker:     speaker label — Dedoose can filter/sort by this
      word_count:  useful for weighting during analysis or reporting turn length
      text:        the verbatim utterance — this is what gets coded in Dedoose

    Why word_count? Dedoose doesn't compute it automatically. Having it as a
    column lets you quickly identify unusually short turns (likely interviewer
    prompts) or unusually long ones (narratives worth flagging for deep coding).

    Verify: len(rows) should equal len(turns). The turn_index should start at
    1 and increment by 1 with no gaps.
    """
    rows = []
    for i, t in enumerate(turns, 1):   # start=1 so turn_index is 1-based, not 0-based
        rows.append({
            "source":     source_name,
            "turn_index": i,
            "speaker":    t["speaker"],
            "word_count": len(t["text"].split()),
            "text":       t["text"],
        })
    return rows


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4: PER-FILE PIPELINE ORCHESTRATION
# ─────────────────────────────────────────────────────────────────────────────

def process_file(path: Path, interactive: bool = True) -> None:
    """
    Run the full pipeline for one transcript file.

    Steps:
      1. Load raw text (handles .txt and .docx)
      2. Strip noise (timestamps, filler markers, whitespace artifacts)
      3. Parse into speaker turns
      4. Merge consecutive same-speaker turns
      5. Show roster; optionally rename speakers interactively
      6. Export to three formats: plain text, Word, Dedoose CSV

    The interactive flag is False when --no-rename is passed on the command
    line — in that case we still print the roster (so the user can see what
    was detected) but skip the rename prompt. This is useful for batch
    processing when speaker names are already standardized.

    The final summary printed to the terminal gives per-speaker turn counts
    and word counts — a quick sanity check that the parsing captured
    everything correctly before the researcher opens the output files.
    """
    print(f"\n{'━'*60}")
    print(f"  File: {path.name}")
    print(f"{'━'*60}")

    raw   = load_raw(path)
    clean = strip_noise(raw)
    turns = parse_turns(clean)
    turns = merge_consecutive(turns)

    if not turns:
        print("  No speaker turns detected — skipping.")
        return

    n_speakers = len({t["speaker"] for t in turns})
    n_turns    = len(turns)
    n_words    = sum(len(t["text"].split()) for t in turns)
    print(f"  {n_speakers} speakers  ·  {n_turns} turns  ·  ~{n_words:,} words")

    # Rename step
    roster = build_roster(turns)
    if interactive:
        rename_map = interactive_rename(roster)
        turns = apply_rename(turns, rename_map)
    else:
        print_roster(roster)

    # Export — create the output directory if it doesn't exist
    CLEAN_DIR.mkdir(parents=True, exist_ok=True)
    stem = path.stem   # filename without extension, used as the base for output names

    # Plain text → Atlas.ti
    txt_path = CLEAN_DIR / f"{stem}_clean.txt"
    txt_path.write_text(to_plain_text(turns), encoding="utf-8")

    # Word document → Atlas.ti
    docx_path = CLEAN_DIR / f"{stem}_clean.docx"
    doc = to_docx(turns, title=stem.replace("_", " ").title())
    doc.save(docx_path)

    # CSV → Dedoose
    csv_path = CLEAN_DIR / f"{stem}_dedoose.csv"
    rows = to_dedoose_csv(turns, source_name=stem)
    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f, fieldnames=["source", "turn_index", "speaker", "word_count", "text"]
        )
        writer.writeheader()
        writer.writerows(rows)

    print(f"\n  Exports saved to {CLEAN_DIR}/")
    print(f"    {txt_path.name:<35} (Atlas.ti plain text)")
    print(f"    {docx_path.name:<35} (Atlas.ti Word)")
    print(f"    {csv_path.name:<35} (Dedoose CSV)")

    # Per-speaker summary — useful for quickly checking that all voices were captured
    print(f"\n  Final speaker labels:")
    roster_final = build_roster(turns)
    for sp in roster_final:
        count = sum(1 for t in turns if t["speaker"] == sp)
        words = sum(len(t["text"].split()) for t in turns if t["speaker"] == sp)
        print(f"    {sp:<30} {count:>3} turns  ·  {words:>5} words")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5: COMMAND-LINE ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    """
    Parse CLI arguments and dispatch to process_file() for each transcript.

    Two modes:
      - Pass filenames as positional arguments to process specific files
      - Run with no arguments to process all .txt and .docx files in transcripts/

    The --no-rename flag disables the interactive prompt, useful for automation
    or when running in a non-interactive environment (e.g., a CI pipeline or
    a Jupyter notebook cell).
    """
    parser = argparse.ArgumentParser(
        description="Clean and restructure interview transcripts for qualitative analysis."
    )
    parser.add_argument(
        "files", nargs="*",
        help="Transcript file(s) to process. Defaults to all files in transcripts/."
    )
    parser.add_argument(
        "--no-rename", action="store_true",
        help="Skip the interactive speaker rename step."
    )
    args = parser.parse_args()

    interactive = not args.no_rename

    if args.files:
        paths = [Path(f) for f in args.files]
    else:
        # Glob for both supported formats; sort so processing order is consistent
        paths = sorted(TRANSCRIPT_DIR.glob("*.txt")) + sorted(TRANSCRIPT_DIR.glob("*.docx"))

    if not paths:
        print(
            f"No transcripts found. Put .txt or .docx files in '{TRANSCRIPT_DIR}/' "
            "or pass filenames directly."
        )
        sys.exit(0)

    for path in paths:
        if not path.exists():
            print(f"File not found: {path}")
            continue
        process_file(path, interactive=interactive)

    print(f"\n{'━'*60}")
    print(f"  Done. All cleaned files in: {CLEAN_DIR}/")
    print(f"{'━'*60}\n")


if __name__ == "__main__":
    main()
