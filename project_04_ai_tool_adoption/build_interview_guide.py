"""
Generates the semi-structured interview guide for Project 4.
Plain Word aesthetic — no colored fills, no design headers.
Looks like something typed into Word on a Tuesday afternoon.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size = 13 if level == 1 else 11
    set_font(run, size=size, bold=True)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def note(text):
    """Inline coaching note — italicized, slightly indented, muted color."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=10, italic=True, color=(100, 100, 100))
    return p

def question(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run_num = p.add_run(f'{num}.  ')
    set_font(run_num, size=11, bold=True)
    run_q = p.add_run(text)
    set_font(run_q, size=11)
    return p

def probe(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    # Remove the default bullet and replace with em-dash
    p.clear()
    p.paragraph_format.left_indent = Inches(0.55)
    run = p.add_run(f'—  {text}')
    set_font(run, size=10.5, color=(80, 80, 80))
    return p

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Semi-Structured Interview Guide')
set_font(run, size=14, bold=True)

p2 = doc.add_paragraph()
run2 = p2.add_run('AI Tool Adoption Study  |  Research Analyst Interviews  |  ~45 minutes')
set_font(run2, size=11, color=(80, 80, 80))

p3 = doc.add_paragraph()
run3 = p3.add_run('Version 1.2  |  For internal use only  |  Do not share with participants')
set_font(run3, size=10, italic=True, color=(120, 120, 120))

rule()

# ═══════════════════════════════════════════════════════════════════════════
# INTERVIEWER NOTES
# ═══════════════════════════════════════════════════════════════════════════
heading('Before You Begin', level=1)
body(
    'Read through this guide the morning of each interview. The goal is not '
    'to get through every question — it is to understand what drove this person\'s '
    'relationship with the tool. Some interviews will spend 20 minutes on Section 3. '
    'That is fine. Follow the participant.'
)
note('[Do not bring a laptop to the interview if possible. A printed guide and a legal pad '
     'signal that you are here to listen, not to document in real time.]')

body(
    'This guide follows a grand tour structure: we start broad (their overall relationship '
    'with research tools) and narrow toward the specific inflection points we care about '
    '(the first time the tool failed, the decision to keep using it or stop). '
    'Do not skip to the specific too early.'
)

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1: INTRO / CONSENT
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 1 — Introduction and Consent', level=1)
note('[~5 minutes. Read this section aloud or closely paraphrase it.]')

body(
    'Thank you for making time today. My name is [interviewer], and I\'m part of the '
    'team evaluating how analysts have been using the new AI platform. This is not a '
    'performance review — there is no right answer. We are trying to understand what '
    'is working, what is not, and why, so we can improve the tool and the onboarding '
    'process before the next rollout.'
)
body(
    'This conversation will take about 45 minutes. I will be taking notes. With your '
    'permission, I would also like to record the audio so I can focus on our conversation '
    'rather than my notes. The recording stays within our team and will not be shared '
    'externally. You can ask me to stop recording at any time.'
)
body('Do I have your permission to record?')
note('[Wait for explicit confirmation. If no: acknowledge, set aside the recorder, continue.]')
body('Do you have any questions before we start?')
note('[Pause. Let silence work. Do not rush past this.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2: KICKOFF / WARM-UP
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 2 — Kickoff', level=1)
note('[~5 minutes. Get them talking. Any answer here is the right answer.]')

question(1, 'Tell me a little about your role — what does a typical research project '
            'look like for you from start to finish?')
note('[This is a grand tour question. Let them take as long as they need. '
     'You are looking for: what tools they already use, where qual vs. quant work sits '
     'in their process, what the pain points are before you even mention the AI tool.]')
probe('What part of that process takes up most of your time?')
probe('What part do you enjoy least?')

question(2, 'Before this platform was introduced, how did you handle qualitative '
            'analysis — transcript coding, open-end processing, that kind of work?')
note('[Listen for: Atlas.ti / NVivo / manual, time spent, frustration level. '
     'This establishes the baseline they are comparing the new tool against.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3: TOOL RELATIONSHIP (CORE)
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 3 — Their Relationship with the Tool', level=1)
note('[~15-20 minutes. This is the core section. Do not rush it.]')

question(3, 'Walk me through the first time you used the platform on a real project — '
            'not the training session, but when it was actually on the line.')
note('[This is the most important question in the guide. "On the line" is intentional — '
     'it signals you want the real story, not the polished version. '
     'Let them tell it without interruption for at least 90 seconds.]')
probe('What were you trying to do?')
probe('What happened?')
probe('What did you do next?')

question(4, 'Describe a moment when the tool gave you something that surprised you — '
            'in either direction.')
note('[You are fishing for trust calibration here. Did they catch an error? '
     'Did they trust an output they probably should not have? '
     'Follow whatever direction they go — do not lead them toward a failure story '
     'if they want to tell a success story first.]')
probe('How did you know whether to trust that output?')
probe('Did that change how you used the tool afterward?')

question(5, 'Tell me about a time the tool did not do what you needed it to do.')
note('[If they say "it always works fine," pause and try: "Even something small — '
     'a moment where you had to step in and fix something, or do it another way." '
     'Everyone has one of these. The goal is not to get them to criticize the tool; '
     'it is to understand how they respond when it fails.]')
probe('What did you do in that moment?')
probe('Did you log it anywhere, or flag it to anyone?')
probe('Did it change how much you used the tool after that?')

question(6, 'How would you describe your relationship with the tool right now — '
            'if you had to put it in a sentence?')
note('[Open-ended. Do not suggest metaphors. Let them find their own language. '
     'Some people say "I use it like a first draft." Some say "I check everything it tells me." '
     'Some say "I stopped using it after week 3." Whatever they say is data.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4: BEHAVIORAL TRACE
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 4 — Usage Patterns Over Time', level=1)
note('[~10 minutes. You are trying to understand the adoption arc — the ramp-up, '
     'the drop, or the plateau — and what caused it.]')

question(7, 'Think back to the first few weeks after onboarding. How often were you '
            'using the platform, compared to now?')
note('[If they say "about the same," probe for why. If they say "a lot less now," '
     'probe for the specific moment things changed — was it one bad experience, '
     'or a slow drift?]')
probe('Was there a specific point when your usage shifted up or down?')
probe('What was happening on your project at that time?')

question(8, 'When you decide to use the tool on a task versus doing it manually — '
            'what goes through your mind?')
note('[This is a decision process question. You want to understand the calculus: '
     'speed vs. accuracy, trust vs. effort, deadline pressure vs. quality concerns. '
     'Do not suggest the categories — let them tell you what matters to them.]')

question(9, 'If a new colleague joined your team tomorrow, would you recommend they '
            'use the platform? What would you tell them?')
note('[The "new colleague" frame lowers stakes and gets you the honest take. '
     'People are more willing to say "I would tell them to be careful about X" '
     'than "I personally distrust X." Pay attention to the caveats they include.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5: ONBOARDING REFLECTION
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 5 — Onboarding and Support', level=1)
note('[~5 minutes. Keep this brief. You are looking for specific gaps, not general ratings.]')

question(10, 'Looking back at your onboarding session, what prepared you well — '
             'and what did you wish had been covered?')
probe('Was there a scenario you ran into later that you did not feel ready for?')
probe('Did you go back to any onboarding materials after the initial session?')

question(11, 'When you hit a problem with the tool, where did you go for help?')
note('[Options include: colleagues, documentation, the support team, or giving up. '
     'Where they go tells you about the informal support network and documentation quality.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6: WRAP-UP
# ═══════════════════════════════════════════════════════════════════════════
heading('Section 6 — Wrap-Up', level=1)
note('[~5 minutes. Wind down the energy. End on a forward-looking note.]')

question(12, 'If you could change one thing about the platform or how it was introduced '
             'to your team, what would it be?')
note('[The most actionable question in the guide. Take verbatim notes here.]')

question(13, 'Is there anything I did not ask about that you think is important for us '
             'to understand?')
note('[This is not a throwaway question. Some of the richest material comes here. '
     'Pause after they respond. If there is a silence, let it breathe for 5 seconds '
     'before you close out.]')

body(
    'That is everything I had. Thank you — this has been genuinely helpful. '
    'The goal is to make sure the next team that goes through this has a better '
    'experience, so I appreciate you being honest with me.'
)
note('[After the recording ends: jot down your immediate impressions while they are fresh. '
     'What was the emotional tone? What surprised you? What did they seem reluctant to say? '
     'Write these field notes before your next interview.]')

rule()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX: QUICK PROBE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════
heading('Appendix — Probe Reference Card', level=1)
note('[Print and keep on your lap during the interview as a fallback.]')

body('When they give you a surface answer:')
probe('Can you tell me more about that?')
probe('What do you mean when you say [repeat their word back]?')
probe('What does that look like in practice?')

body('When they give you a conclusion instead of a story:')
probe('Can you walk me through a specific example?')
probe('Tell me about the last time that happened.')
probe('What were you doing right before that?')

body('When they go quiet:')
probe('[Silence. Wait at least 5 seconds before speaking.]')
probe('Take your time — I want to make sure I understand it correctly.')

body('When they seem to be holding something back:')
probe('Is there anything about that you\'d describe differently if this were off the record?')
probe('What\'s the part that\'s hardest to explain to someone who wasn\'t there?')

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = 'interview_guide_ai_tool_adoption.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
